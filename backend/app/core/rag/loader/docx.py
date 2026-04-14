from __future__ import annotations

import io
from app.core.common.logging import TerraLogUtil as logger
from pathlib import Path
from typing import ClassVar

from app.core.rag.loader.base import BaseLoader
from app.core.rag.models.document import Document, Metadata



class DocxLoader(BaseLoader):
    """DOCX loader using python-docx."""

    supported_extensions: ClassVar[list[str]] = [".docx"]

    def load(self, source: str | Path | bytes, **kwargs) -> list[Document]:
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX loading. Install with: pip install fastrag[docx]")

        raw = self._read_bytes(source)
        source_name = self._resolve_source_name(source)
        doc = docx.Document(io.BytesIO(raw))

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n".join(rows))

        content = "\n\n".join(paragraphs)
        logger.info(
            "Loaded DOCX: source=%s, len=%d",
            source_name, len(text),
        )
        return [
            Document(
                content=content,
                metadata=Metadata(source=source_name, file_type="docx"),
            )
        ]
