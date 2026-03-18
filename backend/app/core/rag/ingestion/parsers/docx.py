"""DOCX file parser."""

from typing import List, Set

from docx import Document as DocxDocument

from llama_index.core.schema import Document

from app.core.rag.ingestion.parsers.base import ParserStrategy


class DocxParser(ParserStrategy):
    """Parse DOCX files using python-docx."""

    @property
    def name(self) -> str:
        return "docx"

    @property
    def supported_extensions(self) -> Set[str]:
        return {"docx"}

    def parse(self, file_path: str) -> List[Document]:
        doc = DocxDocument(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not full_text.strip():
            return []
        return [Document(text=full_text, metadata={"source": file_path})]
