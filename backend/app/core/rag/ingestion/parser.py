"""Document parser using LlamaIndex readers.

Supports: TXT, PDF, DOCX, Markdown, HTML.
Each parser converts a file into a list of LlamaIndex Document objects.
"""

import os
from typing import Dict, List, Type

from llama_index.core.schema import Document

from app.core.common.logging import TerraLogUtil

# Supported file extensions (lowercase, without dot)
SUPPORTED_FILE_TYPES = {"txt", "pdf", "docx",
                        "md", "markdown", "html", "htm", "csv"}


def parse_file(file_path: str, file_type: str) -> List[Document]:
    """Parse a file into LlamaIndex Document objects.

    Uses specialized parsers per file type; falls back to plain-text.

    Args:
        file_path: Absolute path to the file.
        file_type: Extension without dot (e.g. "pdf").

    Returns:
        List of LlamaIndex Document objects.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file type is unsupported.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = file_type.lower().lstrip(".")
    if ext not in SUPPORTED_FILE_TYPES:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {', '.join(sorted(SUPPORTED_FILE_TYPES))}"
        )

    parser = _PARSERS.get(ext, _parse_txt)
    docs = parser(file_path)

    TerraLogUtil.info(
        "file_parsed",
        file_path=file_path,
        file_type=ext,
        doc_count=len(docs),
    )
    return docs


# --------------- per-type parsers ---------------


def _parse_txt(file_path: str) -> List[Document]:
    """Parse a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    if not text.strip():
        return []
    return [Document(text=text, metadata={"source": file_path})]


def _parse_pdf(file_path: str) -> List[Document]:
    """Parse a PDF file using pypdf, one Document per page."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    docs: List[Document] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            docs.append(
                Document(
                    text=text,
                    metadata={"source": file_path, "page": i + 1},
                )
            )
    return docs


def _parse_docx(file_path: str) -> List[Document]:
    """Parse a DOCX file using python-docx."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not full_text.strip():
        return []
    return [Document(text=full_text, metadata={"source": file_path})]


def _parse_markdown(file_path: str) -> List[Document]:
    """Parse a Markdown file as plain text."""
    return _parse_txt(file_path)


def _parse_html(file_path: str) -> List[Document]:
    """Parse an HTML file, stripping tags."""
    from bs4 import BeautifulSoup

    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        html = f.read()

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    if not text.strip():
        return []
    return [Document(text=text, metadata={"source": file_path})]


def _parse_csv(file_path: str) -> List[Document]:
    """Parse a CSV file, treating each row as text."""
    import csv

    rows: List[str] = []
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(", ".join(row))

    if not rows:
        return []
    text = "\n".join(rows)
    return [Document(text=text, metadata={"source": file_path})]


# Registry
_PARSERS: Dict[str, type] = {
    "txt": _parse_txt,
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "md": _parse_markdown,
    "markdown": _parse_markdown,
    "html": _parse_html,
    "htm": _parse_html,
    "csv": _parse_csv,
}
